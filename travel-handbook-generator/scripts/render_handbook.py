#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import subprocess
import textwrap
from io import BytesIO
from pathlib import Path
from urllib.request import Request, urlopen

from PIL import Image, ImageColor, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DEFAULT_THEME = {
    "primary": "0F4C5C",
    "accent": "D48841",
    "soft_blue": "EAF1F6",
    "soft_gold": "FBF3E7",
    "soft_green": "EAF5F3",
    "soft_bg": "F6FBFA",
    "text_dark": "1E2A30",
    "text_muted": "5F7482",
}


def load_spec(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def merged_theme(spec: dict) -> dict:
    out = DEFAULT_THEME.copy()
    out.update(spec.get("theme", {}))
    return out


def rgb(hex_color: str) -> RGBColor:
    value = ImageColor.getrgb(f"#{hex_color}")
    return RGBColor(*value)


def pil_rgb(hex_color: str) -> tuple[int, int, int]:
    return ImageColor.getrgb(f"#{hex_color}")


def load_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in font_candidates:
        if Path(candidate).exists():
            try:
                return ImageFont.truetype(candidate, size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def download_image(url: str, target: Path) -> Path:
    req = Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urlopen(req, timeout=30) as response:
        data = response.read()
    image = Image.open(BytesIO(data)).convert("RGB")
    image.save(target, quality=92)
    return target


def cover_resize(image: Image.Image, width: int, height: int) -> Image.Image:
    source_ratio = image.width / image.height
    target_ratio = width / height
    if source_ratio > target_ratio:
        scale_height = height
        scale_width = int(height * source_ratio)
    else:
        scale_width = width
        scale_height = int(width / source_ratio)
    resized = image.resize((scale_width, scale_height), Image.Resampling.LANCZOS)
    left = (resized.width - width) // 2
    top = (resized.height - height) // 2
    return resized.crop((left, top, left + width, top + height))


def add_overlay(canvas: Image.Image, opacity: int = 118) -> None:
    overlay = Image.new("RGBA", canvas.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)
    draw.rectangle((0, 0, canvas.width, canvas.height), fill=(0, 0, 0, opacity))
    canvas.alpha_composite(overlay)


def resolve_paths(spec_path: Path, spec: dict) -> dict:
    meta = spec["meta"]
    slug = meta["slug"]
    root = spec_path.parents[1] if spec_path.parent.name == "travel_specs" else spec_path.parent
    tmp_dir = root / "tmp" / "docs" / slug
    img_dir = tmp_dir / "images"
    out_dir = Path(meta["output_dir"])
    docx_path = out_dir / meta["docx_filename"]
    pdf_path = out_dir / meta["pdf_filename"]
    preview_prefix = tmp_dir / "preview"
    return {
        "root": root,
        "tmp_dir": tmp_dir,
        "img_dir": img_dir,
        "out_dir": out_dir,
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "preview_prefix": preview_prefix,
    }


def ensure_dirs(paths: dict) -> None:
    paths["img_dir"].mkdir(parents=True, exist_ok=True)
    paths["out_dir"].mkdir(parents=True, exist_ok=True)


def build_cover_image(spec: dict, image_paths: dict[str, Path], paths: dict, theme: dict) -> Path:
    width, height = 1600, 920
    canvas = Image.new("RGBA", (width, height), pil_rgb(theme["soft_bg"]) + (255,))
    hero_keys = spec["cover"]["hero_keys"]
    hero_paths = [image_paths[key] for key in hero_keys]
    slots = [(0, 0, 540, height), (530, 0, 540, height), (1060, 0, 540, height)]
    for index, slot in enumerate(slots):
        source = Image.open(hero_paths[index % len(hero_paths)]).convert("RGB")
        tile = cover_resize(source, slot[2], slot[3])
        canvas.paste(tile, slot[:2])
    add_overlay(canvas, opacity=116)
    draw = ImageDraw.Draw(canvas)
    title_font = load_font(62, bold=True)
    subtitle_font = load_font(30)
    note_font = load_font(22)
    box_top = 470
    draw.rounded_rectangle((80, box_top, width - 80, 840), radius=34, fill=(255, 255, 255, 210), outline=pil_rgb(theme["soft_blue"]), width=3)
    draw.text((120, box_top + 48), spec["meta"]["title"], font=title_font, fill=pil_rgb(theme["primary"]))
    draw.text((120, box_top + 132), spec["meta"]["subtitle"], font=subtitle_font, fill=pil_rgb(theme["text_dark"]))
    draw.text((120, box_top + 200), spec["cover"]["tagline"], font=note_font, fill=pil_rgb(theme["text_muted"]))
    draw.text((120, box_top + 272), f"查询日期：{spec['meta']['query_date']}", font=note_font, fill=pil_rgb(theme["accent"]))
    out = paths["img_dir"] / "cover.png"
    canvas.convert("RGB").save(out, quality=95)
    return out


def build_route_overview_image(spec: dict, paths: dict, theme: dict) -> Path:
    width, height = 1600, 540
    canvas = Image.new("RGB", (width, height), pil_rgb(theme["soft_bg"]))
    draw = ImageDraw.Draw(canvas)
    title_font = load_font(44, bold=True)
    text_font = load_font(24)
    card_title = load_font(28, bold=True)
    small_font = load_font(20)
    draw.rounded_rectangle((48, 40, width - 48, height - 40), radius=28, fill=pil_rgb("FFFFFF"), outline=pil_rgb("D5E6EA"), width=3)
    draw.text((78, 72), "路线总览", font=title_font, fill=pil_rgb(theme["primary"]))
    draw.text((78, 126), spec["overview"]["intro"], font=text_font, fill=pil_rgb(theme["text_muted"]))
    cards = spec["overview"]["route_cards"]
    line_y = 290
    start_x, end_x = 140, width - 140
    draw.line((start_x, line_y, end_x, line_y), fill=pil_rgb("7BAAB4"), width=8)
    step = (end_x - start_x) / max(len(cards) - 1, 1)
    fills = [theme["soft_gold"], theme["soft_blue"], theme["soft_green"], "F1EEF9", "FDF4F1"]
    for idx, card in enumerate(cards):
        x = int(start_x + idx * step)
        draw.ellipse((x - 20, line_y - 20, x + 20, line_y + 20), fill=pil_rgb(theme["accent"]), outline=pil_rgb("FFFFFF"), width=4)
        draw.rounded_rectangle((x - 115, 325, x + 115, 468), radius=24, fill=pil_rgb(fills[idx % len(fills)]), outline=pil_rgb("D7E4E8"), width=2)
        draw.text((x - 78, 345), f"{card['day']}  {card['date']}", font=card_title, fill=pil_rgb(theme["primary"]))
        draw.text((x - 78, 382), card["city"], font=text_font, fill=pil_rgb(theme["accent"]))
        draw.multiline_text((x - 78, 414), textwrap.fill(card["detail"], width=11), font=small_font, fill=pil_rgb(theme["text_dark"]), spacing=5)
    out = paths["img_dir"] / "route_overview.png"
    canvas.save(out, quality=94)
    return out


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def set_default_font(document: Document) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    document.styles["Normal"].font.size = Pt(10.5)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_el = borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            borders.append(edge_el)
        edge_el.set(qn("w:val"), "nil")


def add_page_break(document: Document) -> None:
    document.add_page_break()


def add_title_block(document: Document, title: str, subtitle: str, theme: dict) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(22)
    run.font.color.rgb = rgb(theme["primary"])
    sub = document.add_paragraph(subtitle)
    sub.style = document.styles["Normal"]
    sub.paragraph_format.space_after = Pt(6)
    sub.runs[0].font.color.rgb = rgb(theme["text_muted"])


def add_image(document: Document, path: Path, *, width_cm: float, caption: str | None = None, theme: dict) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = document.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(8.5)
        cp.runs[0].font.color.rgb = rgb(theme["text_muted"])


def add_summary_box(document: Document, title: str, lines: list[str], *, fill: str, theme: dict) -> None:
    fill_map = {
        "blue": theme["soft_blue"],
        "gold": theme["soft_gold"],
        "green": theme["soft_green"],
        "white": "FFFFFF",
    }
    table = document.add_table(rows=1 + len(lines), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    remove_table_borders(table)
    header = table.rows[0].cells[0]
    header.text = title
    header.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_map.get(fill, theme["soft_blue"]))
    header._tc.get_or_add_tcPr().append(shading)
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "PingFang SC"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
            run.font.color.rgb = rgb(theme["primary"])
    for idx, line in enumerate(lines, start=1):
        cell = table.rows[idx].cells[0]
        cell.text = line
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_timeline(document: Document, rows: list[list[str] | tuple[str, str]], theme: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.columns[0].width = Cm(2.7)
    table.columns[1].width = Cm(12.5)
    head_cells = table.rows[0].cells
    head_cells[0].text = "时间"
    head_cells[1].text = "安排"
    for cell in head_cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), theme["primary"])
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = "PingFang SC"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    for time_text, detail in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(time_text)
        row_cells[1].text = str(detail)


def add_hotel_card(document: Document, card_title: str, hotel: dict, image_path: Path, theme: dict) -> None:
    add_title_block(document, card_title, hotel["title"], theme)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Cm(10.8)
    right.width = Cm(4.6)
    left.text = f"当前核实房价：{hotel['price']}\n\n{hotel['note']}\n\n预订链接：{hotel['page']}"
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run().add_picture(str(image_path), width=Cm(4.2))


def add_hotel_options(document: Document, card_title: str, hotel_keys: list[str], spec: dict, theme: dict) -> None:
    add_title_block(document, card_title, "以下为当前日期可选酒店，按当前实时查询结果整理。", theme)
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["酒店", "价格", "选择理由", "链接"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), theme["primary"])
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = "PingFang SC"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    for key in hotel_keys:
        hotel = spec["hotels"][key]
        cells = table.add_row().cells
        cells[0].text = str(hotel["title"])
        cells[1].text = str(hotel["price"])
        cells[2].text = str(hotel["note"])
        cells[3].text = str(hotel["page"])


def add_price_table(document: Document, items: list[list[str] | tuple[str, str, str]], title: str, theme: dict) -> None:
    add_title_block(document, title, "以下金额按当前查询结果整理。", theme)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    heads = ["分类", "说明", "价格"]
    for idx, text in enumerate(heads):
        cell = table.rows[0].cells[idx]
        cell.text = text
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), theme["primary"])
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = "PingFang SC"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    for label, detail, price in items:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(detail)
        cells[2].text = str(price)


def add_source_list(document: Document, sources: list[list[str] | tuple[str, str, str]], query_date: str, theme: dict) -> None:
    add_title_block(document, "来源页", f"以下链接用于核实价格、图像与路线，查询日期为 {query_date}。", theme)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["类别", "说明", "链接"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), theme["primary"])
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for category, label, url in sources:
        cells = table.add_row().cells
        cells[0].text = str(category)
        cells[1].text = str(label)
        cells[2].text = str(url)


def materialize_images(spec: dict, paths: dict) -> dict[str, Path]:
    image_paths: dict[str, Path] = {}
    for key, item in spec["images"].items():
        suffix = ".png"
        image_paths[key] = download_image(item["image"], paths["img_dir"] / f"{key}{suffix}")
    for key, item in spec.get("hotels", {}).items():
        if item.get("image"):
            image_paths[key] = download_image(item["image"], paths["img_dir"] / f"{key}.jpg")
    return image_paths


def build_document(spec: dict, image_paths: dict[str, Path], paths: dict, theme: dict) -> Document:
    document = Document()
    set_page_layout(document)
    set_default_font(document)

    add_image(document, image_paths["cover"], width_cm=16.2, theme=theme)
    add_title_block(document, "行程总览", spec["overview"]["summary"], theme)
    add_image(document, image_paths["route_overview"], width_cm=16.0, caption=spec["overview"]["image_caption"], theme=theme)
    for box in spec["overview"]["summary_boxes"]:
        add_summary_box(document, box["title"], box["lines"], fill=box.get("fill", "blue"), theme=theme)

    for day in spec["days"]:
        add_page_break(document)
        add_title_block(document, day["title"], day["subtitle"], theme)
        image_item = spec["images"][day["image_key"]]
        add_image(document, image_paths[day["image_key"]], width_cm=15.8, caption=image_item["caption"], theme=theme)
        add_timeline(document, day["timeline"], theme)
        for box in day.get("summary_boxes", []):
            add_summary_box(document, box["title"], box["lines"], fill=box.get("fill", "blue"), theme=theme)
        hotel_card = day.get("hotel_card")
        if hotel_card:
            hotel = spec["hotels"][hotel_card["hotel_key"]]
            add_hotel_card(document, hotel_card["card_title"], hotel, image_paths[hotel_card["hotel_key"]], theme)
        hotel_options = day.get("hotel_options")
        if hotel_options:
            add_hotel_options(document, hotel_options["card_title"], hotel_options["hotel_keys"], spec, theme)

    add_page_break(document)
    budget = spec["budget"]
    add_price_table(document, budget["items"], budget["title"], theme)
    total_p = document.add_paragraph()
    total_p.add_run("主预算总计：").bold = True
    total_run = total_p.add_run(budget["total"])
    total_run.bold = True
    total_run.font.color.rgb = rgb(theme["accent"])
    total_p.add_run(budget["note"])
    if budget.get("reference_items"):
        add_price_table(document, budget["reference_items"], budget["reference_title"], theme)
    if budget.get("exclusions"):
        add_summary_box(document, budget["exclusions_title"], budget["exclusions"], fill="gold", theme=theme)

    add_page_break(document)
    add_source_list(document, spec["sources"], spec["meta"]["query_date"], theme)
    return document


def convert_to_pdf(paths: dict) -> None:
    if paths["pdf_path"].exists():
        paths["pdf_path"].unlink()
    subprocess.run(
        [
            "/opt/homebrew/bin/soffice",
            f"-env:UserInstallation=file://{paths['tmp_dir'] / 'lo_profile'}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(paths["out_dir"]),
            str(paths["docx_path"]),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def render_preview(paths: dict) -> None:
    for old in paths["tmp_dir"].glob("preview-*.png"):
        old.unlink()
    subprocess.run(
        [
            "/opt/homebrew/bin/pdftoppm",
            "-png",
            str(paths["pdf_path"]),
            str(paths["preview_prefix"]),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def main() -> None:
    parser = argparse.ArgumentParser(description="Render a DOCX/PDF travel handbook from a JSON spec.")
    parser.add_argument("spec", help="Path to the handbook JSON spec")
    args = parser.parse_args()

    spec_path = Path(args.spec).expanduser().resolve()
    spec = load_spec(spec_path)
    theme = merged_theme(spec)
    paths = resolve_paths(spec_path, spec)
    ensure_dirs(paths)
    image_paths = materialize_images(spec, paths)
    image_paths["cover"] = build_cover_image(spec, image_paths, paths, theme)
    image_paths["route_overview"] = build_route_overview_image(spec, paths, theme)
    document = build_document(spec, image_paths, paths, theme)
    document.save(paths["docx_path"])
    convert_to_pdf(paths)
    render_preview(paths)
    print(f"DOCX: {paths['docx_path']}")
    print(f"PDF: {paths['pdf_path']}")


if __name__ == "__main__":
    main()
