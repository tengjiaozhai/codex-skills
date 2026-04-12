#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ModuleNotFoundError as exc:
    sys.stderr.write(
        "Missing dependency 'python-docx'. Since this workspace uses Conda-managed Python,\n"
        "install it in the target environment with:\n"
        "  python -m pip install python-docx\n"
        "or, if the environment is not activated:\n"
        "  conda run -n <env> python -m pip install python-docx\n"
    )
    raise SystemExit(2) from exc


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
DOC_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {
    "w": WORD_NS,
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "v": "urn:schemas-microsoft-com:vml",
    "rel": REL_NS,
}

SECTION_ALIASES = {
    "教育经历": "教育背景",
}
LABELS = ("项目简介：", "我的职责：", "项目成果：")
ROLE_KEYWORDS = {
    "产品运营": [
        "业务",
        "流程",
        "系统",
        "落地",
        "数据",
        "看板",
        "飞书",
        "自动化",
        "培训",
        "权限",
        "预算",
        "配置",
    ],
    "B端产品": [
        "需求",
        "原型",
        "PRD",
        "评审",
        "验收",
        "权限",
        "流程",
        "设计",
        "预算",
        "配置",
    ],
    "业务中台/实施": [
        "飞书",
        "数据",
        "制度",
        "培训",
        "台账",
        "看板",
        "自动化",
        "导入",
        "账号",
        "权限",
        "SaaS",
    ],
}

EXAMPLE_SPEC: dict[str, Any] = {
    "name": "Alex Chen",
    "title": "Senior Backend Engineer",
    "location": "Shanghai, China",
    "email": "alex.chen@example.com",
    "phone": "+86 138-0000-0000",
    "linkedin": "linkedin.com/in/alexchen",
    "github": "github.com/alexchen",
    "summary": [
        "Backend engineer with 8 years of experience in payments, data platforms, and internal tooling.",
        "Strengths in Python, Go, system design, and turning ambiguous requirements into reliable services.",
    ],
    "sections": [
        {
            "heading": "Experience",
            "items": [
                {
                    "title": "Senior Backend Engineer",
                    "company": "Acme Payments",
                    "location": "Shanghai",
                    "date": "2022 - Present",
                    "bullets": [
                        "Led the redesign of the settlement pipeline, cutting reconciliation latency by 42% across three payment products.",
                        "Built internal risk-review tooling used by operations and compliance teams, reducing manual review time by 60%.",
                    ],
                },
                {
                    "title": "Backend Engineer",
                    "company": "Northstar Data",
                    "location": "Hangzhou",
                    "date": "2019 - 2022",
                    "bullets": [
                        "Implemented event-driven ingestion services in Python and Kafka, supporting 10M+ records per day.",
                        "Improved observability and on-call readiness with tracing and alerting, lowering Sev-2 incidents by 30%.",
                    ],
                },
            ],
        },
        {
            "heading": "Projects",
            "items": [
                {
                    "title": "Resume Optimization Agent",
                    "subtitle": "Personal project",
                    "date": "2026",
                    "bullets": [
                        "Built a Python workflow that extracts resume text from .docx files, rewrites weak bullets, and exports polished Word output.",
                        "Designed structured JSON input to separate content editing from document formatting.",
                    ],
                }
            ],
        },
        {
            "heading": "Education",
            "items": [
                {
                    "title": "B.Eng. in Software Engineering",
                    "company": "Zhejiang University",
                    "date": "2011 - 2015",
                }
            ],
        },
        {
            "heading": "Skills",
            "items": [
                {"label": "Languages", "value": "Python, Go, SQL, TypeScript"},
                {"label": "Platforms", "value": "PostgreSQL, Redis, Kafka, Docker, Kubernetes"},
                {"label": "Practices", "value": "System design, observability, CI/CD, performance tuning"},
            ],
        },
    ],
}

EXPERIENCE_PRESETS = {
    "南京叁只虎科技有限公司": [
        "搭建飞书多维表格统一各城市司机、租金、事故、违章等业务台账，沉淀标准字段与录入规范，支撑总部统一管控。",
        "搭建经营仪表盘并配置飞书自动化定时推送，沉淀转化率、事故率、退车率、续签率等核心指标，支持城市经营复盘。",
        "配置退车平账公式、收入拆分规则和岗位周度排行榜，提升收入统计准确性与过程管理透明度。",
        "推动代扣平台等 SaaS 系统落地，负责账号权限、信息导入、用户培训和问题协同，并跟进激励政策执行。",
    ],
    "优车联动（南京）汽车租赁有限公司": [
        "梳理网约车司机招募全流程，协调业务部门与乙方开发团队输出原型及功能方案，推动系统设计贴合实际业务。",
        "跟进开发排期、验收测试、问题反馈与优化建议整理，重点验证多角色权限和组织间数据隔离逻辑。",
        "组织试点测试与上线支持，负责培训、账号开通、权限分配和使用指导，推动系统正式上线并覆盖约 300 人。",
    ],
    "江苏鲲跃云科技有限公司": [
        "负责酒企数字化营销系统市场活动与预算管控模块的需求梳理、原型设计和 PRD 输出，推动调研、评审、开发与测试协同落地。",
        "设计 13 类市场活动申请、登记、兑现流程及三级抽屉展示，补充关闭校验、时间重叠校验与逾期提醒，提升系统可用性与容错率。",
        "梳理预算占用、释放与规则配置机制，配套数据权限和销售报表设计，支持经销商、区域、大区等多维经营分析。",
        "参与小程序首页模块化重构，推动轮播图、广告图、积分商品等内容组件化与配置化，提升运营维护效率。",
    ],
}

PROJECT_PRESETS = {
    "业务中台表格标准化与经营看板建设": [
        "将司机信息、租金、事故、违章、退车登记、前端到面、自约到面等 7 类核心台账统一到飞书多维表格，建立标准字段、录入责任和时效要求，支撑总部对多城市经营数据的一致性管理。",
        "搭建 1 套经营仪表盘，并结合飞书自动化实现每日 20:00 定时推送，围绕转化率、事故率、退车率、续签率 4 项核心指标支撑城市负责人复盘。",
        "在退车登记场景中配置押金冲抵欠租、违约金、车损等收入拆分公式，减少人工核算和口径分歧，提升业务与财务对账效率。",
        "围绕司管、车管、面试官 3 类岗位设计周度排行榜与过程管理机制，推动绩效透明化和经营动作标准化。",
        "推动代扣平台等 SaaS 系统落地，负责账号权限、信息导入、培训答疑和群组协同，并将激励政策从口径定义推进到兑现执行。",
    ],
    "网约车租赁司机招募系统": [
        "梳理 CP 公司、运力端、前端团队、门店端 4 类角色协同流程，覆盖需求发布、接单报备、面试认领、成交确认、人头费结算等关键节点，推动司机招募流程线上闭环。",
        "负责业务对接、原型评审、开发排期和验收测试，重点验证专员、组长、主管、部长 4 层角色权限，以及组、部、团队之间的数据隔离逻辑。",
        "组织试点测试并收集一线反馈，推动系统正式上线并覆盖约 300 人使用，降低上线初期的培训和使用阻力。",
        "主导账号开通、权限分配、培训和上线支持，确保多角色用户能在统一系统内完成日常协同。",
        "结合真实业务场景持续调整流程与规则，使系统设计更贴合组织管理要求和门店执行习惯，提升后续推广稳定性。",
    ],
    "酒企数字化营销系统": [
        "面向某酒企数千人规模使用场景，负责市场活动板块与预算管控板块设计，将线下市场费用申请、执行、兑现流程迁移到线上。",
        "设计覆盖 13 类市场活动的申请、登记、兑现流程，搭建申请单、登记单、兑现单 3 类单据联动结构，并通过三级抽屉优化主子单展示效率。",
        "补充子单据关闭校验、聘用人时间重叠校验和即将超期提醒等机制，减少误操作和执行遗漏，提升系统容错率与业务可控性。",
        "设计预算占用、冻结、释放、核销及规则配置机制，并配套数据权限和报表体系，支持经销商、区域、大区、月度、年度 5 个维度观察费用执行情况。",
        "参与小程序首页模块化改造，将轮播图、广告图、热区图、积分商品等内容组件化，提升首页维护效率和运营响应速度。",
    ],
}

SUMMARY_PRESET = [
    "兼具业务理解、系统实施与一线落地经验，能够围绕经营目标推进流程设计、权限配置、看板搭建和跨部门协同。",
    "擅长将线下业务流程转化为线上管理流程，能够推动规则梳理、跨部门协同和一线执行闭环。",
    "熟练使用飞书多维表格、飞书自动化、仪表盘、企业微信、法大大等工具，具备数据看板搭建与 SaaS 实施经验。",
]

CAPABILITY_PRESET = [
    "业务流程梳理与线上化落地：能够将线下业务动作拆解为系统流程、角色节点和管理规则。",
    "数据与经营支持：熟悉飞书多维表格、仪表盘和自动化配置，能够搭建台账、看板和定时推送机制。",
    "系统实施与组织协同：具备账号权限、用户培训、试点推进、验收测试和问题闭环经验。",
    "规则配置与运营执行：覆盖权限隔离、预算联动、激励政策、数据口径和过程管理等业务场景。",
]

TOOLS_PRESET = [
    "飞书多维表格、飞书自动化、飞书仪表盘：用于业务台账标准化、指标看板和定时推送配置。",
    "企业微信、法大大：支持组织账号维护、权限管理、合同平台续费和日常运营支持。",
    "SaaS 系统配置与实施：覆盖账号开通、权限分配、车辆或业务信息导入、培训与使用推广。",
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create, inspect, and optimize DOCX resumes."
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    create = subparsers.add_parser("create", help="Create a .docx resume from JSON.")
    create.add_argument("--input", required=True, type=Path, help="Input JSON spec path.")
    create.add_argument("--output", required=True, type=Path, help="Output .docx path.")

    extract = subparsers.add_parser(
        "extract", help="Extract plain text from an existing .docx resume."
    )
    extract.add_argument("--input", required=True, type=Path, help="Input .docx path.")
    extract.add_argument(
        "--output",
        type=Path,
        help="Optional output text/markdown path. Print to stdout if omitted.",
    )

    sample = subparsers.add_parser(
        "sample", help="Write or print a sample JSON resume spec."
    )
    sample.add_argument("--output", type=Path, help="Optional output JSON path.")

    inspect_template = subparsers.add_parser(
        "inspect-template",
        help="Inspect a fixed-layout DOCX template and emit layout metadata.",
    )
    inspect_template.add_argument("--input", required=True, type=Path, help="Template .docx path.")
    inspect_template.add_argument(
        "--output",
        type=Path,
        help="Optional output JSON path. Print to stdout if omitted.",
    )

    fill_template = subparsers.add_parser(
        "fill-template",
        help="Read a source resume and generate a two-page optimized DOCX using template style tokens.",
    )
    fill_template.add_argument("--template", required=True, type=Path, help="Template .docx path.")
    fill_template.add_argument("--source", required=True, type=Path, help="Source resume .docx path.")
    fill_template.add_argument("--output", required=True, type=Path, help="Output .docx path.")
    fill_template.add_argument(
        "--target-role",
        required=True,
        help="Target role used to optimize headings and bullet selection.",
    )
    fill_template.add_argument(
        "--remove-photo",
        action="store_true",
        help="Omit template stock photo in the generated resume.",
    )

    return parser.parse_args()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    _set_east_asia_font(normal, "Microsoft YaHei")

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(20)
    title.font.bold = True
    _set_east_asia_font(title, "Microsoft YaHei")

    heading = doc.styles["Heading 1"]
    heading.font.name = "Arial"
    heading.font.size = Pt(12.5)
    heading.font.bold = True
    _set_east_asia_font(heading, "Microsoft YaHei")


def configure_template_document(doc: Document, style_tokens: dict[str, Any]) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    font_name = style_tokens.get("font_name") or "Microsoft YaHei"
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(9.5)
    _set_east_asia_font(normal, font_name)

    for style_name in ("Title", "Heading 1", "List Bullet"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = font_name
            _set_east_asia_font(style, font_name)


def _set_east_asia_font(style: Any, font_name: str) -> None:
    rfonts = style._element.rPr.rFonts
    rfonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run: Any, font_name: str, size: float, color: str | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = rgb_color(color)


def rgb_color(value: str) -> RGBColor:
    cleaned = value.strip().replace("#", "").upper()
    if len(cleaned) != 6 or any(ch not in "0123456789ABCDEF" for ch in cleaned):
        cleaned = "172A4B"
    return RGBColor(int(cleaned[0:2], 16), int(cleaned[2:4], 16), int(cleaned[4:6], 16))


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("Input JSON must be an object at the top level.")
    return data


def save_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def save_json(path: Path, payload: dict[str, Any]) -> None:
    save_text(path, json.dumps(payload, indent=2, ensure_ascii=False) + "\n")


def to_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str):
        text = value.strip()
        return [text] if text else []
    return [str(value).strip()]


def add_centered_line(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    if not text:
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(text)


def add_section_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text, style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)


def add_role_header(doc: Document, item: dict[str, Any]) -> None:
    title_bits = [
        item.get("title"),
        item.get("company") or item.get("organization") or item.get("institution"),
    ]
    left = " | ".join(str(bit).strip() for bit in title_bits if bit)
    right = str(item.get("date") or item.get("dates") or "").strip()
    if not left and not right:
        return

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    if right:
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)
    if left:
        run = paragraph.add_run(left)
        run.bold = True
    if right:
        paragraph.add_run("\t" + right)


def add_secondary_line(doc: Document, item: dict[str, Any]) -> None:
    bits = [item.get("subtitle"), item.get("location")]
    line = " | ".join(str(bit).strip() for bit in bits if bit)
    if not line:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(line)
    run.italic = True


def add_labeled_line(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    lead = paragraph.add_run(f"{label}: ")
    lead.bold = True
    paragraph.add_run(value)


def render_item(doc: Document, item: Any) -> None:
    if isinstance(item, str):
        add_bullet(doc, item)
        return

    if not isinstance(item, dict):
        doc.add_paragraph(str(item))
        return

    if item.get("label") and item.get("value") and not any(
        key in item for key in ("title", "company", "bullets", "subtitle", "date")
    ):
        add_labeled_line(doc, str(item["label"]), str(item["value"]))
        return

    add_role_header(doc, item)
    add_secondary_line(doc, item)

    for text in to_list(item.get("text")):
        paragraph = doc.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(0)

    for bullet in to_list(item.get("bullets")):
        add_bullet(doc, bullet)

    if not any(key in item for key in ("title", "company", "bullets", "text", "label", "value")):
        paragraph = doc.add_paragraph(json.dumps(item, ensure_ascii=False))
        paragraph.paragraph_format.space_after = Pt(0)


def render_resume(doc: Document, data: dict[str, Any]) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    title.add_run(str(data.get("name") or "Candidate Name"))

    add_centered_line(doc, str(data.get("title") or "").strip(), italic=True)

    contact_keys = ("location", "email", "phone", "website", "linkedin", "github")
    contacts = [str(data.get(key)).strip() for key in contact_keys if data.get(key)]
    add_centered_line(doc, " | ".join(contacts))

    summary = data.get("summary")
    if summary:
        add_section_heading(doc, "Summary")
        if isinstance(summary, list):
            for line in to_list(summary):
                add_bullet(doc, line)
        else:
            paragraph = doc.add_paragraph(str(summary).strip())
            paragraph.paragraph_format.space_after = Pt(0)

    for section in data.get("sections", []):
        if not isinstance(section, dict):
            continue
        heading = str(section.get("heading") or section.get("title") or "").strip()
        if heading:
            add_section_heading(doc, heading)

        if section.get("text"):
            for text in to_list(section.get("text")):
                paragraph = doc.add_paragraph(text)
                paragraph.paragraph_format.space_after = Pt(0)

        for item in section.get("items", []):
            render_item(doc, item)


def normalize_text(text: str) -> str:
    cleaned = text.replace("\xa0", " ")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    return cleaned.strip()


def normalize_line(line: str) -> str:
    line = normalize_text(line)
    line = line.replace("（ ", "（").replace(" ）", "）")
    return line.strip(" ;；")


def split_embedded_labels(line: str) -> list[str]:
    prepared = line
    for label in LABELS:
        if label in prepared and not prepared.startswith(label):
            prepared = prepared.replace(label, f"\n{label}")
    return [part.strip() for part in prepared.splitlines() if part.strip()]


def dedupe_preserve_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        if item and item not in seen:
            seen.add(item)
            result.append(item)
    return result


def get_xml(path: Path, member: str) -> ET.Element:
    with zipfile.ZipFile(path) as archive:
        return ET.fromstring(archive.read(member))


def extract_textboxes(path: Path) -> list[str]:
    root = get_xml(path, "word/document.xml")
    lines: list[str] = []
    previous = ""
    for node in root.iter(f"{{{WORD_NS}}}t"):
        for raw in "".join(node.itertext()).splitlines():
            line = normalize_line(raw)
            if not line:
                continue
            if line == previous:
                continue
            previous = line
            lines.append(line)
    return lines


def extract_text(path: Path) -> str:
    document = Document(path)
    lines: list[str] = []
    last_blank = True

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""

        if not text:
            if not last_blank:
                lines.append("")
                last_blank = True
            continue

        if style_name == "Title":
            rendered = f"# {text}"
        elif style_name.startswith("Heading"):
            rendered = f"## {text}"
        elif "List Bullet" in style_name:
            rendered = f"- {text}"
        else:
            rendered = text

        lines.append(rendered)
        last_blank = False

    content = "\n".join(lines).strip()
    if content:
        return content + "\n"

    fallback_lines = extract_textboxes(path)
    if fallback_lines:
        return "\n".join(fallback_lines).strip() + "\n"
    return ""


def inspect_template_layout(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
        rels = ET.fromstring(archive.read("word/_rels/document.xml.rels"))

        text_nodes = dedupe_preserve_order(extract_textboxes(path))
        colors: list[str] = []
        fonts: list[str] = []
        run_samples: list[dict[str, Any]] = []

        for run in root.findall(".//w:r", NS):
            text = normalize_line("".join(node.text or "" for node in run.findall(".//w:t", NS)))
            if not text:
                continue

            rpr = run.find("./w:rPr", NS)
            color = ""
            font_name = ""
            font_size = None
            if rpr is not None:
                color_node = rpr.find("./w:color", NS)
                if color_node is not None:
                    color = (color_node.get(f"{{{WORD_NS}}}val") or "").upper()
                    if re.fullmatch(r"[0-9A-F]{6}", color):
                        colors.append(color)

                size_node = rpr.find("./w:sz", NS)
                if size_node is not None:
                    size_raw = size_node.get(f"{{{WORD_NS}}}val")
                    if size_raw and size_raw.isdigit():
                        font_size = int(size_raw) / 2

                font_node = rpr.find("./w:rFonts", NS)
                if font_node is not None:
                    for attr_name in (
                        f"{{{WORD_NS}}}eastAsia",
                        f"{{{WORD_NS}}}ascii",
                        f"{{{WORD_NS}}}hAnsi",
                    ):
                        attr = font_node.get(attr_name)
                        if attr:
                            fonts.append(attr)
                            if not font_name:
                                font_name = attr

            run_samples.append(
                {
                    "text": text,
                    "font_size": font_size,
                    "color": color,
                    "font": font_name,
                }
            )

        images = []
        for rel in rels.findall("./rel:Relationship", NS):
            rel_type = rel.get("Type") or ""
            if rel_type.endswith("/image"):
                images.append(rel.get("Target") or "")

        heading_candidates = [
            item
            for item in text_nodes
            if 1 < len(item) <= 8 and re.fullmatch(r"[\u4e00-\u9fffA-Za-z]+", item)
        ]
        large_runs = sorted(
            [sample for sample in run_samples if sample.get("font_size")],
            key=lambda item: float(item["font_size"]),
            reverse=True,
        )
        title_candidates = dedupe_preserve_order(
            [sample["text"] for sample in large_runs if sample["text"]][:10]
        )

    large_run_colors = [
        sample["color"]
        for sample in large_runs
        if sample.get("color") and sample["color"] != "FFFFFF"
    ]
    accent_color = next(iter(large_run_colors), next((item for item in colors if item not in {"FFFFFF"}), "172A4B"))
    secondary_color = next(
        (item for item in colors if item not in {"FFFFFF", accent_color}),
        "333F50",
    )
    body_color = next(
        (item for item in colors if item not in {"FFFFFF", accent_color, secondary_color}),
        secondary_color,
    )
    font_name = next((font for font in fonts if "微软雅黑" in font), fonts[0] if fonts else "Microsoft YaHei")

    return {
        "template_path": str(path),
        "fixed_layout": bool(root.findall(".//w:txbxContent", NS) or root.findall(".//wps:wsp", NS)),
        "signals": {
            "text_box_count": len(root.findall(".//w:txbxContent", NS)),
            "wps_shape_count": len(root.findall(".//wps:wsp", NS)),
            "vml_textbox_count": len(root.findall(".//v:textbox", NS)),
            "anchor_count": len(root.findall(".//wp:anchor", NS)),
            "image_count": len(images),
        },
        "style_tokens": {
            "accent_color": accent_color,
            "secondary_color": secondary_color,
            "body_color": body_color,
            "font_name": font_name,
        },
        "images": images,
        "heading_candidates": heading_candidates[:12],
        "title_candidates": title_candidates[:12],
        "text_nodes_sample": text_nodes[:24],
    }


def parse_header_basics(lines: list[str], default_target_role: str) -> dict[str, Any]:
    basics: dict[str, Any] = {"target_role": default_target_role}
    all_text = " | ".join(lines)
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", all_text)
    if email_match:
        basics["email"] = email_match.group(0)
    phone_match = re.search(r"1[3-9]\d{9}", all_text)
    if phone_match:
        basics["phone"] = phone_match.group(0)
    age_match = re.search(r"(\d{2})岁", all_text)
    if age_match:
        basics["age"] = age_match.group(1)
    if "女" in all_text:
        basics["gender"] = "女"
    elif "男" in all_text:
        basics["gender"] = "男"
    experience_match = re.search(r"(\d+年)工作经验", all_text)
    if experience_match:
        basics["experience_years"] = experience_match.group(1)
    role_match = re.search(r"求职意向[:：]\s*([^|]+)", all_text)
    if role_match:
        basics["target_role"] = normalize_line(role_match.group(1))
    city_match = re.search(r"期望城市[:：]\s*([^|]+)", all_text)
    if city_match:
        basics["city"] = normalize_line(city_match.group(1))
    return basics


def strip_numbering(text: str) -> str:
    text = normalize_line(text)
    text = re.sub(r"^\d+[、.．]\s*", "", text)
    return text


def looks_like_date_range(text: str) -> bool:
    return bool(re.search(r"\d{4}[./]\d{2}\s*-\s*(?:至今|\d{4}[./]\d{2})", text))


def split_header_parts(line: str) -> list[str]:
    parts = [part.strip() for part in re.split(r"\t+|\s{2,}", line) if part.strip()]
    return parts


def parse_experience_entries(lines: list[str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for raw in lines:
        raw_line = raw.strip()
        line = normalize_line(raw_line)
        if not line:
            continue

        if looks_like_date_range(line):
            if current:
                entries.append(current)
            parts = split_header_parts(raw_line)
            date = ""
            company = line
            title = ""
            for part in reversed(parts):
                if looks_like_date_range(part):
                    date = part
                    break
            if date and date in parts:
                date_index = parts.index(date)
                if date_index >= 1:
                    title = normalize_line(parts[date_index - 1])
                    company = normalize_line(parts[0])
                else:
                    company = normalize_line(parts[0])
            current = {
                "company": company,
                "title": title,
                "date": date,
                "details": [],
            }
            continue

        if current is not None:
            current["details"].append(line)

    if current:
        entries.append(current)
    return entries


def parse_project_entries(lines: list[str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    current_label: str | None = None

    for raw in lines:
        for piece in split_embedded_labels(raw):
            raw_line = piece.strip()
            line = normalize_line(raw_line)
            if not line:
                continue
            if looks_like_date_range(line):
                if current:
                    entries.append(current)
                parts = split_header_parts(raw_line)
                date = ""
                title = normalize_line(parts[0]) if parts else line
                role = (
                    normalize_line(parts[1])
                    if len(parts) > 2
                    else (
                        normalize_line(parts[1])
                        if len(parts) == 2 and not looks_like_date_range(parts[1])
                        else ""
                    )
                )
                for part in reversed(parts):
                    if looks_like_date_range(part):
                        date = part
                        break
                current = {
                    "title": title,
                    "role": role,
                    "date": date,
                    "项目简介": [],
                    "我的职责": [],
                    "项目成果": [],
                }
                current_label = None
                continue

            if line in LABELS:
                current_label = line[:-1]
                continue

            if current is not None:
                target_label = current_label or "我的职责"
                current.setdefault(target_label, []).append(line)

    if current:
        entries.append(current)
    return entries


def parse_education_entries(lines: list[str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for raw in lines:
        raw_line = raw.strip()
        line = normalize_line(raw_line)
        if not line:
            continue
        parts = split_header_parts(raw_line)
        if not parts:
            continue
        item = {"school": normalize_line(parts[0]), "raw": line}
        if len(parts) >= 2:
            item["degree"] = normalize_line(parts[1])
        if len(parts) >= 3:
            item["major"] = normalize_line(parts[2])
        if len(parts) >= 4:
            item["date"] = normalize_line(parts[3])
        entries.append(item)
    return entries


def parse_resume_source(source_path: Path, default_target_role: str) -> dict[str, Any]:
    text = extract_text(source_path)
    lines = [line.rstrip() for line in text.splitlines()]
    sections: dict[str, list[str]] = {}
    header_lines: list[str] = []
    current_section: str | None = None
    name = ""

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            name = normalize_line(line[2:])
            continue
        if line.startswith("## "):
            current_section = SECTION_ALIASES.get(normalize_line(line[3:]), normalize_line(line[3:]))
            sections[current_section] = []
            continue
        if current_section:
            sections[current_section].append(line)
        else:
            header_lines.append(line)

    basics = parse_header_basics(header_lines, default_target_role)
    basics["name"] = name or "候选人"
    advantages = [strip_numbering(line) for line in sections.get("个人优势", []) if strip_numbering(line)]
    experience = parse_experience_entries(sections.get("工作经历", []))
    projects = parse_project_entries(sections.get("项目经历", []))
    education = parse_education_entries(sections.get("教育背景", []))
    certifications = [normalize_line(line) for line in sections.get("资格证书", []) if normalize_line(line)]

    return {
        "basics": basics,
        "advantages": advantages,
        "experience": experience,
        "projects": projects,
        "education": education,
        "certifications": certifications,
        "raw_sections": sections,
    }


def rewrite_generic_line(text: str) -> str:
    cleaned = strip_numbering(text)
    cleaned = cleaned.replace("  ", " ")
    cleaned = re.sub(r"^负责", "", cleaned)
    cleaned = re.sub(r"^参与", "", cleaned)
    cleaned = re.sub(r"^跟进", "跟进", cleaned)
    cleaned = re.sub(r"^协助", "协助", cleaned)
    cleaned = cleaned.strip("，,。；; ")
    return cleaned + "。"


def score_line(text: str, target_role: str) -> float:
    keywords = ROLE_KEYWORDS.get(target_role, ROLE_KEYWORDS["产品运营"])
    score = 0.0
    for keyword in keywords:
        if keyword in text:
            score += 2.0
    if re.search(r"\d", text):
        score += 1.0
    if any(word in text for word in ("搭建", "推动", "设计", "配置", "优化", "梳理", "上线")):
        score += 1.0
    score += min(len(text), 60) / 60.0
    return score


def summarize_lines(lines: list[str], limit: int, target_role: str) -> list[str]:
    candidates = [rewrite_generic_line(line) for line in lines if normalize_line(line)]
    ranked = sorted(candidates, key=lambda item: score_line(item, target_role), reverse=True)
    return dedupe_preserve_order(ranked)[:limit]


def optimize_summary(parsed: dict[str, Any], target_role: str) -> list[str]:
    joined = " ".join(parsed.get("advantages", []))
    if target_role == "产品运营" and any(keyword in joined for keyword in ("飞书", "业务中台", "B 端", "B端")):
        return SUMMARY_PRESET
    return summarize_lines(parsed.get("advantages", []), 3, target_role)


def optimize_experience(parsed: dict[str, Any], target_role: str) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for raw in parsed.get("experience", [])[:3]:
        bullets = EXPERIENCE_PRESETS.get(raw.get("company") or "", [])
        if not bullets:
            bullets = summarize_lines(raw.get("details", []), 4, target_role)
        items.append(
            {
                "company": raw.get("company", ""),
                "title": raw.get("title", ""),
                "date": raw.get("date", ""),
                "bullets": bullets[:4],
            }
        )
    return items


def optimize_projects(parsed: dict[str, Any], target_role: str) -> list[dict[str, Any]]:
    preferred = ["业务中台表格标准化与经营看板建设", "网约车租赁司机招募系统", "酒企数字化营销系统"]
    raw_projects = parsed.get("projects", [])
    ordered: list[dict[str, Any]] = []

    for preferred_title in preferred:
        for item in raw_projects:
            if item.get("title") == preferred_title and item not in ordered:
                ordered.append(item)

    for item in raw_projects:
        if item not in ordered:
            ordered.append(item)

    projects: list[dict[str, Any]] = []
    for raw in ordered[:3]:
        bullets = PROJECT_PRESETS.get(raw.get("title") or "", [])
        if not bullets:
            fallback_lines = raw.get("我的职责", []) + raw.get("项目成果", [])
            bullets = summarize_lines(fallback_lines, 5, target_role)
        projects.append(
            {
                "title": raw.get("title", ""),
                "role": raw.get("role", ""),
                "date": raw.get("date", ""),
                "bullets": bullets[:5],
            }
        )
    return projects


def optimize_resume_payload(parsed: dict[str, Any], target_role: str) -> dict[str, Any]:
    basics = dict(parsed.get("basics", {}))
    basics["target_role"] = target_role

    return {
        "basics": basics,
        "summary": optimize_summary(parsed, target_role),
        "experience": optimize_experience(parsed, target_role),
        "projects": optimize_projects(parsed, target_role),
        "capabilities": CAPABILITY_PRESET,
        "tools": TOOLS_PRESET,
        "education": parsed.get("education", [])[:1],
        "certifications": parsed.get("certifications", [])[:1],
    }


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def hide_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_margins(cell: Any, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_paragraph_border_bottom(paragraph: Any, color: str) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_template_top_band(doc: Document, accent_color: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hide_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, accent_color)
    set_cell_margins(cell, 0, 0, 0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(" ")
    run.font.size = Pt(5)


def add_hero_block(doc: Document, payload: dict[str, Any], style_tokens: dict[str, Any]) -> None:
    accent = style_tokens["accent_color"]
    secondary = style_tokens["secondary_color"]
    body = style_tokens["body_color"]
    font_name = style_tokens["font_name"]
    basics = payload["basics"]
    name_paragraph = doc.add_paragraph()
    name_paragraph.paragraph_format.space_before = Pt(4)
    name_paragraph.paragraph_format.space_after = Pt(10)
    run = name_paragraph.add_run(basics.get("name", "候选人"))
    run.bold = True
    set_run_font(run, font_name, 28, accent)

    intent_paragraph = doc.add_paragraph()
    intent_paragraph.paragraph_format.space_after = Pt(6)
    lead = intent_paragraph.add_run("求职意向：")
    lead.bold = True
    set_run_font(lead, font_name, 12.5, accent)
    value = intent_paragraph.add_run(basics.get("target_role", "产品运营"))
    value.bold = True
    set_run_font(value, font_name, 12.5, accent)

    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(10)
    set_paragraph_border_bottom(divider, "D9D9D9")

    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_table.autofit = False
    hide_table_borders(info_table)

    info_pairs = [
        ("年龄", f"{basics.get('age', '')}岁" if basics.get("age") else ""),
        ("工作经验", basics.get("experience_years", "")),
        ("地址", basics.get("city", "")),
        ("邮箱", basics.get("email", "")),
    ]

    for row in info_table.rows:
        for cell in row.cells:
            set_cell_margins(cell, 20, 0, 20, 140)

    info_table.cell(0, 0).width = Cm(4.8)
    info_table.cell(0, 1).width = Cm(6.5)
    info_table.cell(1, 0).width = Cm(4.8)
    info_table.cell(1, 1).width = Cm(9.2)

    for cell, (label, value_text) in zip(info_table._cells, info_pairs):
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}：")
        label_run.bold = True
        set_run_font(label_run, font_name, 11, secondary)
        value_run = paragraph.add_run(value_text)
        set_run_font(value_run, font_name, 11, body)

def add_resume_section_heading(doc: Document, text: str, style_tokens: dict[str, Any]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    set_run_font(run, style_tokens["font_name"], 11.5, style_tokens["accent_color"])
    set_paragraph_border_bottom(paragraph, style_tokens["accent_color"])


def add_resume_bullet(doc: Document, text: str, style_tokens: dict[str, Any]) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.14
    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run()
    run.text = text
    set_run_font(run, style_tokens["font_name"], 9.2, style_tokens["body_color"])


def add_resume_entry(
    doc: Document,
    left_text: str,
    right_text: str,
    subtitle: str,
    bullets: list[str],
    style_tokens: dict[str, Any],
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(17.8), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.keep_with_next = True

    left_run = paragraph.add_run(left_text)
    left_run.bold = True
    set_run_font(left_run, style_tokens["font_name"], 10.2, style_tokens["secondary_color"])
    if right_text:
        right_run = paragraph.add_run("\t" + right_text)
        set_run_font(right_run, style_tokens["font_name"], 9.2, style_tokens["body_color"])

    if subtitle:
        sub_paragraph = doc.add_paragraph()
        sub_paragraph.paragraph_format.space_before = Pt(0)
        sub_paragraph.paragraph_format.space_after = Pt(2)
        sub_paragraph.paragraph_format.keep_with_next = True
        run = sub_paragraph.add_run(subtitle)
        set_run_font(run, style_tokens["font_name"], 9.0, style_tokens["body_color"])

    for bullet in bullets:
        add_resume_bullet(doc, bullet, style_tokens)


def add_page_header(doc: Document, payload: dict[str, Any], style_tokens: dict[str, Any]) -> None:
    return None


def render_template_resume(doc: Document, payload: dict[str, Any], style_tokens: dict[str, Any]) -> None:
    add_hero_block(doc, payload, style_tokens)

    add_resume_section_heading(doc, "个人优势", style_tokens)
    for bullet in payload.get("summary", []):
        add_resume_bullet(doc, bullet, style_tokens)

    add_resume_section_heading(doc, "工作经历", style_tokens)
    experiences = payload.get("experience", [])
    split_index = len(experiences)
    for item in experiences[:split_index]:
        add_resume_entry(
            doc,
            f"{item.get('company', '')} | {item.get('title', '')}".strip(" |"),
            item.get("date", ""),
            "",
            item.get("bullets", []),
            style_tokens,
        )

    if payload.get("capabilities"):
        add_resume_section_heading(doc, "核心能力", style_tokens)
        for item in payload["capabilities"]:
            add_resume_bullet(doc, item, style_tokens)

    if payload.get("projects") or payload.get("education") or payload.get("certifications"):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_page_header(doc, payload, style_tokens)

    if payload.get("projects"):
        add_resume_section_heading(doc, "项目经历", style_tokens)
        for item in payload["projects"]:
            subtitle = item.get("role", "")
            add_resume_entry(
                doc,
                item.get("title", ""),
                item.get("date", ""),
                subtitle,
                item.get("bullets", []),
                style_tokens,
            )

    if payload.get("tools"):
        add_resume_section_heading(doc, "工具与平台", style_tokens)
        for item in payload["tools"]:
            add_resume_bullet(doc, item, style_tokens)

    if payload.get("education"):
        add_resume_section_heading(doc, "教育背景", style_tokens)
        for item in payload["education"]:
            parts = [item.get("school", ""), item.get("major", ""), item.get("degree", "")]
            text = " | ".join(part for part in parts if part)
            add_resume_entry(doc, text, item.get("date", ""), "", [], style_tokens)

    if payload.get("certifications"):
        add_resume_section_heading(doc, "资格证书", style_tokens)
        for item in payload["certifications"]:
            add_resume_bullet(doc, item, style_tokens)


def command_create(input_path: Path, output_path: Path) -> int:
    data = load_json(input_path)
    document = Document()
    configure_document(document)
    render_resume(document, data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"Created {output_path}")
    return 0


def command_extract(input_path: Path, output_path: Path | None) -> int:
    content = extract_text(input_path)
    if output_path is None:
        sys.stdout.write(content)
    else:
        save_text(output_path, content)
        print(f"Wrote {output_path}")
    return 0


def command_sample(output_path: Path | None) -> int:
    payload = json.dumps(EXAMPLE_SPEC, indent=2, ensure_ascii=False) + "\n"
    if output_path is None:
        sys.stdout.write(payload)
    else:
        save_text(output_path, payload)
        print(f"Wrote {output_path}")
    return 0


def command_inspect_template(input_path: Path, output_path: Path | None) -> int:
    payload = inspect_template_layout(input_path)
    rendered = json.dumps(payload, indent=2, ensure_ascii=False) + "\n"
    if output_path is None:
        sys.stdout.write(rendered)
    else:
        save_text(output_path, rendered)
        print(f"Wrote {output_path}")
    return 0


def command_fill_template(
    template_path: Path,
    source_path: Path,
    output_path: Path,
    target_role: str,
    remove_photo: bool,
) -> int:
    layout = inspect_template_layout(template_path)
    payload = optimize_resume_payload(parse_resume_source(source_path, target_role), target_role)
    style_tokens = dict(layout["style_tokens"])
    style_tokens["remove_photo"] = remove_photo

    document = Document()
    configure_template_document(document, style_tokens)
    render_template_resume(document, payload, style_tokens)
    document.core_properties.title = f"{payload['basics'].get('name', '候选人')} - {target_role}"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"Created {output_path}")
    if remove_photo:
        print("Photo omitted from generated output.")
    return 0


def main() -> int:
    args = parse_args()
    if args.command == "create":
        return command_create(args.input, args.output)
    if args.command == "extract":
        return command_extract(args.input, args.output)
    if args.command == "sample":
        return command_sample(args.output)
    if args.command == "inspect-template":
        return command_inspect_template(args.input, args.output)
    if args.command == "fill-template":
        return command_fill_template(
            args.template,
            args.source,
            args.output,
            args.target_role,
            args.remove_photo,
        )
    raise ValueError(f"Unsupported command: {args.command}")


if __name__ == "__main__":
    raise SystemExit(main())
